"""macOS document inserter using Python-docx and clipboard."""

import os
import subprocess
import time
from typing import Optional, List
from ..base import DocumentInserter, SpreadsheetInserter, DocumentInserterFactory

try:
    from docx import Document
except ImportError:
    Document = None


class MacOSWordInserter(DocumentInserter):
    """macOS Word document inserter using python-docx and clipboard."""

    def insert(self, docx_path: str, move_cursor_to_end: bool = True) -> None:
        """
        在 Microsoft Word 中插入 DOCX 文件内容
        
        策略：使用 python-docx 读取 DOCX 内容，通过剪贴板粘贴
        
        Args:
            docx_path: DOCX 文件路径
            move_cursor_to_end: 是否将光标移到插入内容末尾（暂未支持）
            
        Raises:
            FileNotFoundError: 文件不存在
            RuntimeError: 插入失败
        """
        from ...utils.app_logging import log
        log(f"[MacOSWordInserter.insert] START with docx_path={docx_path}")
        
        if not os.path.exists(docx_path):
            raise FileNotFoundError(f"DOCX file not found: {docx_path}")
        
        if Document is None:
            raise RuntimeError("python-docx is not installed")
        
        # 保存原剪贴板内容
        log("[MacOSWordInserter] Saving original clipboard content...")
        try:
            original_clipboard_process = subprocess.Popen(
                ['pbpaste'],
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE
            )
            original_clipboard, _ = original_clipboard_process.communicate()
        except Exception as e:
            log(f"[MacOSWordInserter] Warning: Failed to save clipboard: {e}")
            original_clipboard = None
        
        try:
            # 1. 读取 DOCX 文件
            log("[MacOSWordInserter] Reading DOCX file...")
            doc = Document(docx_path)
            
            # 2. 提取所有段落和表格内容（使用 Document API 而非直接 XML）
            log("[MacOSWordInserter] Extracting content...")
            content_lines = []
            
            for para in doc.paragraphs:
                content_lines.append(para.text)
            
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    content_lines.append('\t'.join(cells))
            
            content = '\n'.join(content_lines)
            log(f"[MacOSWordInserter] Extracted content length: {len(content)}")
            
            # 3. 将内容放到剪贴板
            log("[MacOSWordInserter] Copying to clipboard...")
            process = subprocess.Popen(
                ['pbcopy'],
                stdin=subprocess.PIPE,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE
            )
            process.communicate(content.encode('utf-8'))
            log("[MacOSWordInserter] Content copied to clipboard")
            
            # 4. 使用 Word 的 AppleScript 粘贴命令（而不是键盘快捷键）
            log("[MacOSWordInserter] Delaying 0.3s before paste...")
            time.sleep(0.3)
            
            script = '''
            tell application "Microsoft Word"
                activate
                delay 0.2
                tell active document
                    paste
                end tell
            end tell
            '''
            
            log("[MacOSWordInserter] *** EXECUTING AppleScript paste command ***")
            result = subprocess.run(
                ['osascript', '-e', script],
                capture_output=True,
                text=True,
                timeout=5
            )
            log("[MacOSWordInserter] *** COMPLETED AppleScript paste command ***")
            
            # 再延迟一下确保粘贴完成
            log("[MacOSWordInserter] Delaying 0.1s after paste...")
            time.sleep(0.1)
            
            if result.returncode != 0:
                raise RuntimeError(f"Failed to paste content: {result.stderr}")
            
            log(f"[MacOSWordInserter.insert] COMPLETED successfully")
                
        except Exception as e:
            log(f"[MacOSWordInserter] FAILED with error: {e}")
            if isinstance(e, RuntimeError):
                raise
            raise RuntimeError(f"Failed to insert document: {str(e)}")
        finally:
            # 恢复原剪贴板内容
            if original_clipboard is not None:
                log("[MacOSWordInserter] Restoring original clipboard content...")
                try:
                    restore_process = subprocess.Popen(
                        ['pbcopy'],
                        stdin=subprocess.PIPE,
                        stdout=subprocess.PIPE,
                        stderr=subprocess.PIPE
                    )
                    restore_process.communicate(original_clipboard)
                    log("[MacOSWordInserter] Original clipboard content restored")
                except Exception as e:
                    log(f"[MacOSWordInserter] Warning: Failed to restore clipboard: {e}")

    def is_available(self) -> bool:
        """检查 Microsoft Word 是否可用"""
        script = '''
        tell application "System Events"
            return (name of processes) contains "Microsoft Word"
        end tell
        '''
        result = subprocess.run(
            ['osascript', '-e', script],
            capture_output=True,
            text=True,
            timeout=2
        )
        return result.returncode == 0 and result.stdout.strip().lower() == "true"


class MacOSExcelInserter(SpreadsheetInserter):
    """macOS Excel spreadsheet inserter using TSV format and clipboard."""

    def insert(self, table_data: List[List[str]], keep_format: bool = True) -> bool:
        """
        在 Microsoft Excel 中插入表格数据
        
        策略：将表格数据转换为 TSV 格式，通过剪贴板粘贴到 Excel
        
        Args:
            table_data: 二维数组表格数据
            keep_format: 是否保留格式
            
        Returns:
            True 如果成功插入
        """
        from ...utils.app_logging import log
        
        if not table_data:
            return False
        
        # 保存原剪贴板内容
        log("[MacOSExcelInserter] Saving original clipboard content...")
        try:
            original_clipboard_process = subprocess.Popen(
                ['pbpaste'],
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE
            )
            original_clipboard, _ = original_clipboard_process.communicate()
        except Exception as e:
            log(f"[MacOSExcelInserter] Warning: Failed to save clipboard: {e}")
            original_clipboard = None
        
        try:
            # 1. 将表格转换为 TSV 格式（Excel 能识别）
            log("[MacOSExcelInserter] Converting table to TSV...")
            tsv_lines = []
            for row in table_data:
                # 转义特殊字符
                escaped_cells = [str(cell).replace('\t', '    ').replace('\n', ' ') for cell in row]
                tsv_lines.append('\t'.join(escaped_cells))
            
            tsv_content = '\n'.join(tsv_lines)
            log(f"[MacOSExcelInserter] TSV content length: {len(tsv_content)}")
            
            # 2. 将 TSV 内容放到剪贴板
            log("[MacOSExcelInserter] Copying to clipboard...")
            process = subprocess.Popen(
                ['pbcopy'],
                stdin=subprocess.PIPE,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE
            )
            process.communicate(tsv_content.encode('utf-8'))
            log("[MacOSExcelInserter] Content copied to clipboard")
            
            # 3. 使用 Excel 的 AppleScript 粘贴命令
            log("[MacOSExcelInserter] Delaying before paste...")
            time.sleep(0.3)
            
            script = '''
            tell application "Microsoft Excel"
                activate
                delay 0.2
                tell active sheet
                    paste
                end tell
            end tell
            '''
            
            log("[MacOSExcelInserter] *** EXECUTING AppleScript paste command ***")
            result = subprocess.run(
                ['osascript', '-e', script],
                capture_output=True,
                text=True,
                timeout=5
            )
            log("[MacOSExcelInserter] *** COMPLETED AppleScript paste command ***")
            
            time.sleep(0.1)
            return result.returncode == 0
            
        except Exception as e:
            log(f"[MacOSExcelInserter] FAILED with error: {e}")
            return False
        finally:
            # 恢复原剪贴板内容
            if original_clipboard is not None:
                log("[MacOSExcelInserter] Restoring original clipboard content...")
                try:
                    restore_process = subprocess.Popen(
                        ['pbcopy'],
                        stdin=subprocess.PIPE,
                        stdout=subprocess.PIPE,
                        stderr=subprocess.PIPE
                    )
                    restore_process.communicate(original_clipboard)
                    log("[MacOSExcelInserter] Original clipboard content restored")
                except Exception as e:
                    log(f"[MacOSExcelInserter] Warning: Failed to restore clipboard: {e}")
    
    def insert_table(self, xlsx_path: str) -> None:
        """
        兼容性方法：在 Microsoft Excel 中插入表格数据
        注意：此方法在 macOS 上不支持从文件插入
        
        Args:
            xlsx_path: XLSX 文件路径（未使用）
        """
        raise NotImplementedError("macOS Excel table insertion requires raw table data via insert() method, not file path")

    def is_available(self) -> bool:
        """检查 Microsoft Excel 是否可用"""
        script = '''
        tell application "System Events"
            return (name of processes) contains "Microsoft Excel"
        end tell
        '''
        result = subprocess.run(
            ['osascript', '-e', script],
            capture_output=True,
            text=True,
            timeout=2
        )
        return result.returncode == 0 and result.stdout.strip().lower() == "true"


class MacOSNumbersInserter(SpreadsheetInserter):
    """macOS Numbers spreadsheet inserter."""

    def insert_table(self, xlsx_path: str) -> None:
        """在 Apple Numbers 中插入表格（暂未实现）"""
        raise NotImplementedError("macOS Numbers insertion not yet implemented")

    def is_available(self) -> bool:
        """检查 Apple Numbers 是否可用"""
        script = '''
        tell application "System Events"
            return (name of processes) contains "Numbers"
        end tell
        '''
        result = subprocess.run(
            ['osascript', '-e', script],
            capture_output=True,
            text=True,
            timeout=2
        )
        return result.returncode == 0 and result.stdout.strip().lower() == "true"


class MacOSDocumentInserterFactory(DocumentInserterFactory):
    """Factory for macOS document inserters."""

    def __init__(self):
        self._word_inserter = MacOSWordInserter()
        self._excel_inserter = MacOSExcelInserter()
        self._numbers_inserter = MacOSNumbersInserter()

    def get_word_inserter(self) -> DocumentInserter:
        """Get Word document inserter."""
        return self._word_inserter

    def get_excel_inserter(self) -> SpreadsheetInserter:
        """Get Excel spreadsheet inserter."""
        return self._excel_inserter

    def get_inserter_for_app(self, app_id: str) -> Optional[object]:
        """Get appropriate inserter for detected application."""
        if app_id == "word":
            return self._word_inserter
        elif app_id == "excel":
            return self._excel_inserter
        elif app_id == "numbers":
            return self._numbers_inserter
        else:
            return None


def get_document_inserter_factory():
    return MacOSDocumentInserterFactory()


__all__ = [
    "MacOSWordInserter",
    "MacOSExcelInserter",
    "MacOSNumbersInserter",
    "MacOSDocumentInserterFactory",
    "get_document_inserter_factory",
]